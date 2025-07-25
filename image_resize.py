import cv2
import streamlit as st
import numpy as np
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches
import os

def detect_display_region(image):
    img_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    candidates = []

    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        aspect_ratio = w / float(h)
        area = cv2.contourArea(cnt)
        if 1000 < area < 50000 and 1.5 < aspect_ratio < 6:
            candidates.append((x, y, w, h))

    if candidates:
        x, y, w, h = sorted(candidates, key=lambda b: b[2]*b[3], reverse=True)[0]
        cropped = img_cv[y:y+h, x:x+w]
        cropped_pil = Image.fromarray(cv2.cvtColor(cropped, cv2.COLOR_BGR2RGB))
        return cropped_pil
    return None

def generate_word_table(images):
    doc = Document()
    doc.add_heading('Meter Display Regions', level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Image Name'
    hdr_cells[1].text = 'Display Crop'

    for name, image in images:
        row = table.add_row().cells
        row[0].text = name

        img_stream = BytesIO()
        image.save(img_stream, format='PNG')
        img_stream.seek(0)
        run = row[1].paragraphs[0].add_run()
        run.add_picture(img_stream, width=Inches(2.5))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 🚀 Streamlit UI
st.title("Meter Display Detector (Batch Mode) 📦")

uploaded_files = st.file_uploader("Upload meter images 📷", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

if uploaded_files:
    cropped_images = []

    for uploaded_file in uploaded_files:
        st.subheader(f"🖼️ {uploaded_file.name}")
        image = Image.open(uploaded_file)
        st.image(image, caption="Original", use_column_width=True)

        display_crop = detect_display_region(image)
        if display_crop:
            st.image(display_crop, caption="Detected Display", use_column_width=True)
            cropped_images.append((uploaded_file.name, display_crop))
        else:
            st.warning(f"No display region found in {uploaded_file.name}")

    if cropped_images:
        word_buffer = generate_word_table(cropped_images)

        st.download_button(
            label="📥 Download Word File",
            data=word_buffer,
            file_name="meter_display_table.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
